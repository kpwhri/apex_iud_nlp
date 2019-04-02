import glob
import os
import random
import sys
from collections import defaultdict

from cronkd.util.docx import add_table_of_contents
from docx import Document


def snippet_samples(fp, sample=25, encoding='utf8', pattern='text_*.out'):
    res = defaultdict(lambda: defaultdict(set))
    fn = sorted(glob.glob(os.path.join(fp, pattern)))[-1]
    with open(fn, encoding=encoding) as fh:
        for i, line in enumerate(fh):
            if i == 0:
                continue
            noteid, algo, cat, cat_num, terms, text = line.split('\t')
            res[algo][cat].add(text)

    doc = Document()
    add_table_of_contents(doc)
    for algo, cats in res.items():
        doc.add_heading(algo, level=1)
        for cat, s in cats.items():
            doc.add_heading(cat, level=2)
            for example in random.sample(s, min(sample, len(s))):
                doc.add_paragraph(example, style='List Number')
    doc.save(f'{fn}.docx')


if __name__ == '__main__':
    snippet_samples(sys.argv[1])
